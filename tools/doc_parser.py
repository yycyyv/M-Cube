from __future__ import annotations

import os
import hashlib
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

import fitz  # PyMuPDF
from agents.base_agent import BaseStructuredAgent, RetryPolicy
from docx import Document
from pydantic import BaseModel, ConfigDict, Field
from services.llm_factory import build_llm_callable

from models.image_schemas import ImageAsset


class ParsedDocument(BaseModel):
    """Normalized parser output for downstream agent/tool processing."""

    model_config = ConfigDict(extra="forbid")

    source_path: str = Field(..., min_length=1)
    file_type: str = Field(..., min_length=1, description="pdf/docx/txt")
    text: str = Field(..., min_length=1, description="Full extracted text content.")
    page_count: int | None = Field(default=None, ge=1)
    page_image_paths: list[str] = Field(
        default_factory=list,
        description="Rendered PDF page image paths for OCR fallback and downstream vision.",
    )
    metadata: dict[str, str] = Field(default_factory=dict)


class LLMOCRResult(BaseModel):
    """Structured output schema for vision-LLM OCR fallback."""

    model_config = ConfigDict(extra="forbid")
    text: str = Field(default="", description="Transcribed text only.")


class DocumentParser:
    """
    Basic parser framework.
    Supports .pdf/.doc/.docx/.txt and returns normalized ParsedDocument.
    """

    SUPPORTED_SUFFIXES = {".pdf", ".doc", ".docx", ".txt"}

    def __init__(
        self,
        image_root_dir: str | Path | None = None,
        *,
        llm_runtime: dict[str, str | None] | None = None,
        llm_api_key: str | None = None,
    ) -> None:
        upload_root = Path(os.getenv("UPLOAD_ROOT_DIR", ".runtime/uploads")).resolve()
        self._image_root = Path(image_root_dir).resolve() if image_root_dir else (upload_root / "images")
        self._image_root.mkdir(parents=True, exist_ok=True)
        # Media constraints protect runtime memory/disk and keep image payloads bounded.
        # Keep a sane default for patent disclosures that commonly embed >20 figures.
        self._max_images_per_file = int(os.getenv("UPLOAD_MAX_EMBEDDED_IMAGES", "40"))
        self._max_image_bytes = int(os.getenv("UPLOAD_MAX_IMAGE_BYTES", str(5 * 1024 * 1024)))
        self._max_total_image_bytes = int(os.getenv("UPLOAD_MAX_TOTAL_IMAGE_BYTES", str(25 * 1024 * 1024)))
        self._max_image_pixels = int(os.getenv("UPLOAD_MAX_IMAGE_PIXELS", str(16_000_000)))
        # Heuristics to keep only primary patent-related figures and skip tiny decorative assets.
        self._min_primary_image_side = int(os.getenv("UPLOAD_MIN_PRIMARY_IMAGE_SIDE", "120"))
        self._min_primary_image_area = int(os.getenv("UPLOAD_MIN_PRIMARY_IMAGE_AREA", "30000"))
        self._max_images_per_pdf_page = int(os.getenv("UPLOAD_MAX_IMAGES_PER_PDF_PAGE", "4"))
        self._pdf_text_min_chars_for_non_scan = int(os.getenv("UPLOAD_PDF_TEXT_MIN_CHARS", "50"))
        self._rendered_pdf_dpi = int(os.getenv("UPLOAD_RENDERED_PDF_DPI", "144"))
        self._last_pdf_ocr_error: str | None = None
        runtime = llm_runtime or {}
        self._llm_provider = (runtime.get("provider") or "").strip() or None
        self._llm_model = (runtime.get("model") or "").strip() or None
        self._llm_vision_model = (runtime.get("vision_model") or "").strip() or None
        self._llm_base_url = (runtime.get("base_url") or "").strip() or None
        self._llm_api_key = (llm_api_key or "").strip() or None
        self._llm_ocr_enabled = os.getenv("UPLOAD_LLM_OCR_FALLBACK", "1") != "0"
        self._llm_ocr_agent: BaseStructuredAgent[LLMOCRResult] | None = self._build_llm_ocr_agent()

    def parse_file(self, file_path: str | Path) -> ParsedDocument:
        path = Path(file_path).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(f"Unsupported file type: {suffix}")

        if suffix == ".pdf":
            text, page_count, page_image_paths = self._parse_pdf(path)
            return ParsedDocument(
                source_path=str(path),
                file_type="pdf",
                text=text,
                page_count=page_count,
                page_image_paths=page_image_paths,
                metadata={"filename": path.name},
            )

        if suffix == ".docx":
            text = self._parse_docx(path)
            return ParsedDocument(
                source_path=str(path),
                file_type="docx",
                text=text,
                metadata={"filename": path.name},
            )

        if suffix == ".doc":
            text = self._parse_doc(path)
            return ParsedDocument(
                source_path=str(path),
                file_type="doc",
                text=text,
                metadata={"filename": path.name},
            )

        text = self._parse_txt(path)
        return ParsedDocument(
            source_path=str(path),
            file_type="txt",
            text=text,
            metadata={"filename": path.name},
        )

    def extract_images(self, file_path: str | Path, source_file_id: str = "unknown") -> list[ImageAsset]:
        """
        Extract embedded images from source documents and persist them under runtime image dir.
        Returns image metadata list for downstream workflow state.
        """
        path = Path(file_path).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"Input file not found: {path}")

        suffix = path.suffix.lower()
        if suffix == ".pdf":
            images = self.extract_images_from_pdf(path, source_file_id=source_file_id)
        elif suffix == ".docx":
            images = self.extract_images_from_docx(path, source_file_id=source_file_id)
        elif suffix == ".doc":
            images = self._extract_images_from_doc(path, source_file_id=source_file_id)
        elif suffix == ".txt":
            images = []
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        self._validate_image_batch(images)
        return images

    def extract_images_from_pdf(self, path: Path, source_file_id: str = "unknown") -> list[ImageAsset]:
        """Extract embedded raster images from PDF and persist to runtime image store."""
        doc = fitz.open(path)
        extracted: list[ImageAsset] = []
        seen_hashes: set[str] = set()
        try:
            for page_index in range(doc.page_count):
                page = doc.load_page(page_index)
                page_kept = 0
                for image_info in page.get_images(full=True):
                    if page_kept >= self._max_images_per_pdf_page:
                        break
                    xref = int(image_info[0])
                    image_dict = doc.extract_image(xref)
                    raw = image_dict.get("image")
                    ext = str(image_dict.get("ext", "png")).lower()
                    if not isinstance(raw, (bytes, bytearray)) or len(raw) == 0:
                        continue
                    digest = hashlib.sha1(bytes(raw)).hexdigest()
                    if digest in seen_hashes:
                        continue
                    mime_type = self._ext_to_mime(ext)
                    if not self._is_supported_raster_mime(mime_type):
                        # Skip vector/unknown formats (e.g. jbig2/emf proxies) that vision models cannot open.
                        continue
                    width, height = self._read_dimensions(bytes(raw), mime_type)
                    if not self._is_primary_candidate(width=width, height=height):
                        continue
                    saved_path = self._persist_image_bytes(
                        source_file_id=source_file_id,
                        ext=ext,
                        raw=bytes(raw),
                    )
                    extracted.append(
                        ImageAsset(
                            image_id=Path(saved_path).stem,
                            source_file_id=source_file_id,
                            source_path=saved_path,
                            page_index=page_index,
                            mime_type=mime_type,
                            width=width,
                            height=height,
                            caption_hint=f"{path.name}#page={page_index + 1}",
                        )
                    )
                    seen_hashes.add(digest)
                    page_kept += 1
        finally:
            doc.close()
        return extracted

    def extract_images_from_docx(self, path: Path, source_file_id: str = "unknown") -> list[ImageAsset]:
        """Extract embedded images from DOCX package (word/media/*) and persist to runtime store."""
        extracted: list[ImageAsset] = []
        seen_hashes: set[str] = set()
        try:
            with ZipFile(path, "r") as archive:
                media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
                for name in media_files:
                    raw = archive.read(name)
                    if not raw:
                        continue
                    digest = hashlib.sha1(raw).hexdigest()
                    if digest in seen_hashes:
                        continue
                    ext = Path(name).suffix.lower().lstrip(".") or "png"
                    mime_type = self._ext_to_mime(ext)
                    if not self._is_supported_raster_mime(mime_type):
                        # Skip unsupported media such as emf/wmf in Word templates.
                        continue
                    width, height = self._read_dimensions(raw, mime_type)
                    if not self._is_primary_candidate(width=width, height=height):
                        continue
                    saved_path = self._persist_image_bytes(
                        source_file_id=source_file_id,
                        ext=ext,
                        raw=raw,
                    )
                    extracted.append(
                        ImageAsset(
                            image_id=Path(saved_path).stem,
                            source_file_id=source_file_id,
                            source_path=saved_path,
                            page_index=None,
                            mime_type=mime_type,
                            width=width,
                            height=height,
                            caption_hint=f"{path.name}:{Path(name).name}",
                        )
                    )
                    seen_hashes.add(digest)
        except BadZipFile as exc:
            raise ValueError(f"Failed to parse DOCX images: {exc}") from exc
        return extracted

    def _parse_pdf(self, path: Path) -> tuple[str, int, list[str]]:
        """Extract text from PDF, with OCR fallback for scan-like files."""
        doc = fitz.open(path)
        try:
            pages: list[str] = []
            page_image_paths = self._render_pdf_pages_to_temp(doc=doc, source_name=path.stem)
            for page_index in range(doc.page_count):
                page = doc.load_page(page_index)
                page_text = self._extract_text_from_pdf_page(page).strip()
                if page_text:
                    pages.append(page_text)
            page_count = doc.page_count
            joined_text = "\n\n".join(pages).strip()
            if len(joined_text) < self._pdf_text_min_chars_for_non_scan:
                llm_ocr_text = self._run_pdf_llm_ocr_fallback(page_image_paths=page_image_paths).strip()
                if llm_ocr_text:
                    joined_text = llm_ocr_text
            if not joined_text:
                # Scanned/image-only PDFs are common in OA uploads. Return a stable
                # fallback marker so downstream OA flow can continue with visual data.
                ocr_note = ""
                if self._last_pdf_ocr_error:
                    ocr_note = f"; ocr_status={self._last_pdf_ocr_error}"
                joined_text = (
                    f"[NO_EXTRACTABLE_TEXT_IN_PDF] filename={path.name}; pages={page_count}; "
                    f"This PDF likely contains scanned images only{ocr_note}."
                )
            return joined_text, page_count, page_image_paths
        finally:
            doc.close()

    def _build_llm_ocr_agent(self) -> BaseStructuredAgent[LLMOCRResult] | None:
        if not self._llm_ocr_enabled:
            return None

        provider = (self._llm_provider or "").strip() or os.getenv("LLM_PROVIDER", "").strip()
        runtime_openai_key = self._llm_api_key if provider == "openai" else None
        if not provider and (runtime_openai_key or os.getenv("OPENAI_API_KEY", "").strip()):
            provider = "openai"

        text_model = (
            (self._llm_model or "").strip()
            or os.getenv("LLM_MODEL", "").strip()
            or os.getenv("OPENAI_MODEL", "").strip()
            or os.getenv("OPENAI_VISION_MODEL", "").strip()
        )
        vision_model = (
            (self._llm_vision_model or "").strip()
            or os.getenv("LLM_VISION_MODEL", "").strip()
            or os.getenv("OPENAI_VISION_MODEL", "").strip()
        )
        base_url = (
            (self._llm_base_url or "").strip()
            or os.getenv("LLM_BASE_URL", "").strip()
            or os.getenv("OPENAI_BASE_URL", "").strip()
        )

        llm_callable = build_llm_callable(
            provider=provider or None,
            model=text_model or None,
            vision_model=vision_model or None,
            base_url=base_url or None,
            api_key=self._llm_api_key,
        )
        if llm_callable is None:
            return None
        return BaseStructuredAgent[LLMOCRResult](
            name="pdf_llm_ocr_agent",
            llm_callable=llm_callable,
            retry_policy=RetryPolicy(max_retries=1),
        )

    def _run_pdf_llm_ocr_fallback(self, *, page_image_paths: list[str]) -> str:
        """
        LLM-based OCR fallback:
        send each rendered page to configured vision model and transcribe text only.
        """
        if not self._llm_ocr_enabled or self._llm_ocr_agent is None:
            if self._last_pdf_ocr_error is None:
                self._last_pdf_ocr_error = "llm_ocr_unavailable"
            return ""

        prompt = (
            "你是一个专业的专利 OCR 引擎。"
            "请将这张专利扫描件图片中的所有文字一字不差地转录下来。"
            "不要描述图片结构，只输出文字。"
        )
        lines: list[str] = []
        for image_path in page_image_paths:
            try:
                result = self._llm_ocr_agent.run_structured(
                    prompt=prompt,
                    output_model=LLMOCRResult,
                    context={"image_paths": [image_path], "image_mime_types": ["image/png"]},
                )
                text = result.text.strip()
                if text:
                    lines.append(text)
            except Exception as exc:  # noqa: BLE001 - page-level failures should not break full parsing.
                self._last_pdf_ocr_error = f"llm_ocr_failed: {exc}"
                continue
        return "\n".join(lines).strip()

    def _render_pdf_pages_to_temp(self, *, doc: fitz.Document, source_name: str) -> list[str]:
        """Render every PDF page to PNG under runtime temp image dir and return file paths."""
        tmp_dir = self._image_root / f"pdf_pages_{source_name}_{uuid4().hex}"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        scale = max(1.0, float(self._rendered_pdf_dpi) / 72.0)
        matrix = fitz.Matrix(scale, scale)
        paths: list[str] = []
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            out_path = (tmp_dir / f"page_{page_index + 1:04d}.png").resolve()
            pix.save(str(out_path))
            paths.append(str(out_path))
        return paths

    @staticmethod
    def _extract_text_from_pdf_page(page: fitz.Page) -> str:
        """
        Robust per-page text extraction for mixed PDF generators.
        Try multiple representations to reduce false negatives.
        """
        pieces: list[str] = []

        direct_text = page.get_text("text").strip()
        if direct_text:
            pieces.append(direct_text)

        if not pieces:
            blocks = page.get_text("blocks")
            block_texts = [str(block[4]).strip() for block in blocks if len(block) >= 5 and str(block[4]).strip()]
            if block_texts:
                pieces.append("\n".join(block_texts))

        if not pieces:
            words = page.get_text("words")
            if words:
                # words tuple: x0, y0, x1, y1, word, block_no, line_no, word_no
                words_sorted = sorted(words, key=lambda item: (item[5], item[6], item[7]))
                word_text = " ".join(str(item[4]).strip() for item in words_sorted if str(item[4]).strip())
                if word_text.strip():
                    pieces.append(word_text.strip())

        if not pieces:
            rawdict = page.get_text("rawdict")
            line_chunks: list[str] = []
            for block in rawdict.get("blocks", []):
                for line in block.get("lines", []):
                    span_chunks: list[str] = []
                    for span in line.get("spans", []):
                        # Some PDFs store glyphs only in raw chars without span.text.
                        span_text = str(span.get("text", "")).strip()
                        if span_text:
                            span_chunks.append(span_text)
                            continue
                        chars = span.get("chars", [])
                        if isinstance(chars, list) and chars:
                            rebuilt = "".join(
                                str(ch.get("c", "")) for ch in chars if isinstance(ch, dict) and ch.get("c")
                            ).strip()
                            if rebuilt:
                                span_chunks.append(rebuilt)
                    line_text = "".join(span_chunks).strip()
                    if line_text:
                        line_chunks.append(line_text)
            if line_chunks:
                pieces.append("\n".join(line_chunks))

        return "\n".join(pieces).strip()

    def _parse_docx(self, path: Path) -> str:
        """Extract text from DOCX paragraphs and tables; many templates store body content in tables."""
        doc = Document(str(path))
        collected: list[str] = []
        seen: set[str] = set()

        def _append_piece(piece: str) -> None:
            normalized = piece.replace("\u3000", " ").strip()
            if not normalized:
                return
            # Avoid repeated boilerplate chunks from merged cells/headers.
            if normalized in seen:
                return
            seen.add(normalized)
            collected.append(normalized)

        for paragraph in doc.paragraphs:
            _append_piece(paragraph.text)

        # Some patent disclosure templates put almost all content inside tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _append_piece(cell.text)

        text = "\n".join(collected).strip()
        if not text:
            raise ValueError(f"No extractable text found in DOCX: {path}")
        return text

    def _parse_txt(self, path: Path) -> str:
        """Read UTF-8 text files; fallback errors are surfaced to caller."""
        text = path.read_text(encoding="utf-8").strip()
        if not text:
            raise ValueError(f"Input TXT is empty: {path}")
        return text

    def _parse_doc(self, path: Path) -> str:
        """
        Parse legacy DOC on Windows via Microsoft Word COM automation.
        DOC has no reliable pure-python parser, so we use:
        1) Word COM on Windows (preferred)
        2) LibreOffice soffice headless conversion (fallback)
        """
        com_error_message: str | None = None
        # First try Word COM on Windows if available.
        if os.name == "nt":
            try:
                import pythoncom  # type: ignore[import-untyped]
                import win32com.client  # type: ignore[import-untyped]
            except Exception:  # noqa: BLE001
                win32com = None
            else:
                word = None
                doc = None
                com_initialized = False
                with tempfile.TemporaryDirectory() as tmp_dir:
                    converted = Path(tmp_dir) / f"{path.stem}.docx"
                    try:
                        # COM must be initialized in the current thread before DispatchEx.
                        pythoncom.CoInitialize()
                        com_initialized = True
                        # Launch isolated Word instance for conversion.
                        word = win32com.client.DispatchEx("Word.Application")
                        word.Visible = False
                        word.DisplayAlerts = 0
                        doc = word.Documents.Open(str(path))
                        # 16 = wdFormatDocumentDefault (.docx)
                        doc.SaveAs2(str(converted), FileFormat=16)
                        return self._parse_docx(converted)
                    except Exception as exc:  # noqa: BLE001
                        com_error_message = str(exc)
                    finally:
                        if doc is not None:
                            try:
                                doc.Close(False)
                            except Exception:  # noqa: BLE001
                                pass
                        if word is not None:
                            try:
                                word.Quit()
                            except Exception:  # noqa: BLE001
                                pass
                        if com_initialized:
                            try:
                                pythoncom.CoUninitialize()
                            except Exception:  # noqa: BLE001
                                pass

        # Fallback: try LibreOffice headless conversion if installed.
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_dir = Path(tmp_dir)
            command = [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(path),
            ]
            try:
                subprocess.run(command, check=True, capture_output=True, text=True)
            except FileNotFoundError as exc:
                if com_error_message:
                    raise ValueError(
                        f"DOC parse failed via Microsoft Word COM: {com_error_message}; "
                        "LibreOffice (soffice) is also unavailable."
                    ) from exc
                raise ValueError(
                    "Failed to parse DOC. Install Microsoft Word or LibreOffice (soffice), and ensure pywin32 is installed."
                ) from exc
            except subprocess.CalledProcessError as exc:
                if com_error_message:
                    raise ValueError(
                        f"DOC parse failed via Microsoft Word COM: {com_error_message}; "
                        f"and via LibreOffice: {exc.stderr.strip() or 'conversion error'}"
                    ) from exc
                raise ValueError(
                    f"Failed to parse DOC via LibreOffice: {exc.stderr.strip() or 'conversion error'}"
                ) from exc

            converted = output_dir / f"{path.stem}.docx"
            if not converted.exists():
                raise ValueError("Failed to parse DOC: conversion output not found.")
            return self._parse_docx(converted)

    def _extract_images_from_doc(self, path: Path, source_file_id: str) -> list[ImageAsset]:
        """Convert DOC to DOCX then reuse DOCX image extraction pipeline."""
        com_error_message: str | None = None
        if os.name == "nt":
            try:
                import pythoncom  # type: ignore[import-untyped]
                import win32com.client  # type: ignore[import-untyped]
            except Exception:  # noqa: BLE001
                win32com = None
            else:
                word = None
                doc = None
                com_initialized = False
                with tempfile.TemporaryDirectory() as tmp_dir:
                    converted = Path(tmp_dir) / f"{path.stem}.docx"
                    try:
                        # COM must be initialized in the current thread before DispatchEx.
                        pythoncom.CoInitialize()
                        com_initialized = True
                        word = win32com.client.DispatchEx("Word.Application")
                        word.Visible = False
                        word.DisplayAlerts = 0
                        doc = word.Documents.Open(str(path))
                        doc.SaveAs2(str(converted), FileFormat=16)
                        return self.extract_images_from_docx(converted, source_file_id=source_file_id)
                    except Exception as exc:  # noqa: BLE001
                        com_error_message = str(exc)
                    finally:
                        if doc is not None:
                            try:
                                doc.Close(False)
                            except Exception:  # noqa: BLE001
                                pass
                        if word is not None:
                            try:
                                word.Quit()
                            except Exception:  # noqa: BLE001
                                pass
                        if com_initialized:
                            try:
                                pythoncom.CoUninitialize()
                            except Exception:  # noqa: BLE001
                                pass

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_dir = Path(tmp_dir)
            command = [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(path),
            ]
            try:
                subprocess.run(command, check=True, capture_output=True, text=True)
            except FileNotFoundError as exc:
                if com_error_message:
                    raise ValueError(
                        f"DOC image parse failed via Microsoft Word COM: {com_error_message}; "
                        "LibreOffice (soffice) is also unavailable."
                    ) from exc
                raise ValueError(
                    "Failed to parse DOC images. Install Microsoft Word or LibreOffice (soffice), "
                    "and ensure pywin32 is installed."
                ) from exc
            except subprocess.CalledProcessError as exc:
                if com_error_message:
                    raise ValueError(
                        f"DOC image parse failed via Microsoft Word COM: {com_error_message}; "
                        f"and via LibreOffice: {exc.stderr.strip() or 'conversion error'}"
                    ) from exc
                raise ValueError(
                    f"Failed to parse DOC images via LibreOffice: {exc.stderr.strip() or 'conversion error'}"
                ) from exc

            converted = output_dir / f"{path.stem}.docx"
            if not converted.exists():
                raise ValueError("Failed to parse DOC images: conversion output not found.")
            return self.extract_images_from_docx(converted, source_file_id=source_file_id)

    def _persist_image_bytes(self, *, source_file_id: str, ext: str, raw: bytes) -> str:
        ext_clean = ext.lstrip(".").lower() or "png"
        image_id = f"{source_file_id}_{uuid4()}"
        target = (self._image_root / f"{image_id}.{ext_clean}").resolve()
        target.write_bytes(raw)
        return str(target)

    def _validate_image_batch(self, images: list[ImageAsset]) -> None:
        if len(images) > self._max_images_per_file:
            raise ValueError(
                "Embedded image count exceeds limit: "
                f"{len(images)} > {self._max_images_per_file}."
            )
        total_bytes = 0
        for image in images:
            image_path = Path(image.source_path)
            try:
                size = image_path.stat().st_size
            except OSError:
                size = 0
            total_bytes += size
            if size > self._max_image_bytes:
                raise ValueError(
                    "Embedded image size exceeds limit: "
                    f"{size} > {self._max_image_bytes} bytes ({image_path.name})."
                )
            if image.width and image.height and (image.width * image.height) > self._max_image_pixels:
                raise ValueError(
                    "Embedded image pixel area exceeds limit: "
                    f"{image.width}x{image.height} > {self._max_image_pixels} pixels ({image_path.name})."
                )
        if total_bytes > self._max_total_image_bytes:
            raise ValueError(
                "Total embedded image bytes exceed limit: "
                f"{total_bytes} > {self._max_total_image_bytes}."
            )

    @staticmethod
    def _ext_to_mime(ext: str) -> str:
        mapping = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "webp": "image/webp",
            "tif": "image/tiff",
            "tiff": "image/tiff",
        }
        return mapping.get(ext.lower(), "application/octet-stream")

    @staticmethod
    def _is_supported_raster_mime(mime_type: str) -> bool:
        return mime_type in {
            "image/png",
            "image/jpeg",
            "image/gif",
            "image/bmp",
            "image/webp",
            "image/tiff",
        }

    def _is_primary_candidate(self, *, width: int | None, height: int | None) -> bool:
        # Unknown dimensions are treated as non-primary to avoid noisy binary assets.
        if width is None or height is None:
            return False
        if width < self._min_primary_image_side or height < self._min_primary_image_side:
            return False
        if (width * height) < self._min_primary_image_area:
            return False
        return True

    @staticmethod
    def _read_dimensions(raw: bytes, mime_type: str) -> tuple[int | None, int | None]:
        # Best-effort width/height parsing without extra heavy dependencies.
        if len(raw) < 10:
            return None, None
        if mime_type == "image/png" and raw.startswith(b"\x89PNG\r\n\x1a\n") and len(raw) >= 24:
            width = int.from_bytes(raw[16:20], "big")
            height = int.from_bytes(raw[20:24], "big")
            return width or None, height or None
        if mime_type == "image/gif" and (raw.startswith(b"GIF87a") or raw.startswith(b"GIF89a")):
            width = int.from_bytes(raw[6:8], "little")
            height = int.from_bytes(raw[8:10], "little")
            return width or None, height or None
        if mime_type == "image/bmp" and raw.startswith(b"BM") and len(raw) >= 26:
            width = int.from_bytes(raw[18:22], "little", signed=True)
            height = int.from_bytes(raw[22:26], "little", signed=True)
            return abs(width) or None, abs(height) or None
        if mime_type == "image/jpeg" and raw.startswith(b"\xff\xd8"):
            stream = BytesIO(raw)
            stream.read(2)
            while True:
                marker_prefix = stream.read(1)
                if not marker_prefix:
                    break
                if marker_prefix != b"\xff":
                    continue
                marker = stream.read(1)
                while marker == b"\xff":
                    marker = stream.read(1)
                if not marker:
                    break
                marker_int = marker[0]
                if marker_int in {0xD8, 0xD9}:
                    continue
                length_bytes = stream.read(2)
                if len(length_bytes) != 2:
                    break
                segment_len = int.from_bytes(length_bytes, "big")
                if segment_len < 2:
                    break
                if marker_int in {0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF}:
                    payload = stream.read(segment_len - 2)
                    if len(payload) >= 5:
                        height = int.from_bytes(payload[1:3], "big")
                        width = int.from_bytes(payload[3:5], "big")
                        return width or None, height or None
                    break
                stream.seek(segment_len - 2, os.SEEK_CUR)
        return None, None

